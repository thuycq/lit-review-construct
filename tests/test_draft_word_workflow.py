import json
from pathlib import Path

from docx import Document

from litreview_construct.draft_support import prepare_working_draft_packet, save_working_draft
from litreview_construct.project import init_project
from litreview_construct.researcher_package import prepare_researcher_package
from litreview_construct.word_export import export_artifact_docx
from litreview_construct.workflow import project_next_step


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload), encoding="utf-8")


def _write_jsonl(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("".join(json.dumps(row) + "\n" for row in rows), encoding="utf-8")


def _accepted_blueprint_fixture(root: Path) -> None:
    init_project(root, name="Draft Test")
    state_path = root / ".litreview" / "state.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    state["stages"]["literature_review_blueprint"]["status"] = "accepted"
    state["stages"]["researcher_handoff"]["status"] = "in_progress"
    state_path.write_text(json.dumps(state), encoding="utf-8")

    _write_json(
        root / ".litreview" / "data" / "selected_direction.json",
        {"title": "Asymmetric WCM underinvestment", "supporting_paper_ids": ["p1", "p2"]},
    )
    _write_json(
        root / ".litreview" / "data" / "blueprint.json",
        {
            "schema_version": 1,
            "revision": 2,
            "title": "Working-capital literature review",
            "organizing_logic": "Move from nonlinear optimality to SME financing constraints.",
            "sections": [
                {
                    "section_id": "s1",
                    "order": 1,
                    "title": "Nonlinear optimality",
                    "purpose": "Establish the benchmark logic.",
                    "key_arguments": ["Both too little and too much working capital may be costly."],
                    "anchor_paper_ids": ["p1"],
                    "supporting_paper_ids": [],
                    "conflicting_paper_ids": [],
                    "evidence_ids": ["e1"],
                    "theoretical_foundations": [],
                    "methodological_context": [],
                    "hypothesis_or_proposition_links": ["H1"],
                    "unresolved_questions": [],
                    "transition_logic": "Move to SME asymmetry.",
                },
                {
                    "section_id": "s2",
                    "order": 2,
                    "title": "Financing constraints",
                    "purpose": "Explain why underinvestment may be more severe for SMEs.",
                    "key_arguments": ["Financial constraints may intensify underinvestment."],
                    "anchor_paper_ids": ["p2"],
                    "supporting_paper_ids": [],
                    "conflicting_paper_ids": [],
                    "evidence_ids": ["e2"],
                    "theoretical_foundations": [],
                    "methodological_context": [],
                    "hypothesis_or_proposition_links": ["H2"],
                    "unresolved_questions": [],
                    "transition_logic": None,
                },
            ],
            "opening_tasks": [],
            "cross_section_synthesis_tasks": [],
            "closing_tasks": [],
            "verification_priorities": [],
            "limitations": [],
        },
    )
    _write_jsonl(
        root / ".litreview" / "data" / "papers.jsonl",
        [
            {"paper_id": "p1", "title": "Nonlinear WCM", "year": 2022},
            {"paper_id": "p2", "title": "SME finance constraints", "year": 2023, "file_reference": ".litreview/cache/fulltext/p2.pdf"},
        ],
    )
    _write_jsonl(
        root / ".litreview" / "data" / "evidence.jsonl",
        [
            {
                "evidence_id": "e1",
                "paper_id": "p1",
                "evidence_type": "association",
                "claim": "The abstract reports a nonlinear association.",
                "provenance": "source_reported",
                "source_basis": "abstract",
                "certainty": "medium",
            },
            {
                "evidence_id": "e2",
                "paper_id": "p2",
                "evidence_type": "association",
                "claim": "The full text links financing constraints to working-capital policy.",
                "provenance": "source_reported",
                "source_basis": "full_text",
                "certainty": "high",
            },
        ],
    )


def test_working_draft_flags_abstract_evidence(tmp_path: Path) -> None:
    _accepted_blueprint_fixture(tmp_path)
    packet = prepare_working_draft_packet(tmp_path)
    assert packet["sections"] == 2
    assert packet["abstract_only_evidence"] == 1

    submission = {
        "title": "Researcher working draft",
        "sections": [
            {
                "section_id": "s1",
                "title": "Nonlinear optimality",
                "fragments": [
                    {
                        "purpose": "Establish benchmark logic",
                        "draft_text": "Working-capital policy is commonly framed as a balancing problem in which both shortages and excesses may be associated with weaker performance.",
                        "paper_ids": ["p1"],
                        "evidence_ids": ["e1"],
                        "researcher_tasks": ["Verify the exact functional form."],
                        "verification_notes": [],
                    }
                ],
                "researcher_decisions": [],
            },
            {
                "section_id": "s2",
                "title": "Financing constraints",
                "fragments": [
                    {
                        "purpose": "Explain the SME mechanism",
                        "draft_text": "Financing constraints provide a plausible mechanism through which SMEs may operate below their feasible working-capital benchmark.",
                        "paper_ids": ["p2"],
                        "evidence_ids": ["e2"],
                        "researcher_tasks": [],
                        "verification_notes": [],
                    }
                ],
                "researcher_decisions": [],
            },
        ],
        "cross_section_notes": [],
        "final_researcher_tasks": [],
        "limitations": ["One fragment remains abstract-grounded."],
    }
    input_file = tmp_path / ".litreview" / "packets" / "working_draft_submission.json"
    _write_json(input_file, submission)
    saved = save_working_draft(tmp_path, input_file)
    assert saved["verification_fragments"] == 1
    text = (tmp_path / "outputs" / "06b_literature_review_working_draft.md").read_text(encoding="utf-8")
    assert "VERIFY BEFORE USE" in text
    assert "Researcher Working Draft" in text


def test_word_export_creates_editable_docx(tmp_path: Path) -> None:
    init_project(tmp_path, name="Word Test")
    output_md = tmp_path / "outputs" / "06b_literature_review_working_draft.md"
    output_md.parent.mkdir(parents=True, exist_ok=True)
    output_md.write_text(
        "# Researcher Working Draft — Literature Review\n\n## 1. Nonlinear optimality\n\nA researcher-editable paragraph.\n\n| Evidence | Basis |\n|---|---|\n| e1 | abstract |\n",
        encoding="utf-8",
    )
    result = export_artifact_docx(tmp_path, artifact="working-draft")
    path = Path(result["output"])
    assert path.is_file()
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Researcher Working Draft" in text
    assert "Nonlinear optimality" in text
    assert len(document.tables) == 1


def _workflow_fixture(root: Path, *, with_evidence: bool, with_working_draft: bool) -> None:
    init_project(root, name="Workflow Test")
    state_path = root / ".litreview" / "state.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    for stage in ("research_intent", "seed_literature", "research_direction", "literature_review_blueprint"):
        state["stages"][stage]["status"] = "accepted"
    state["stages"]["literature_discovery"]["status"] = "ready_for_review"
    state["stages"]["evidence_mapping"]["status"] = "ready_for_review" if with_evidence else "not_started"
    state_path.write_text(json.dumps(state), encoding="utf-8")
    _write_json(root / ".litreview" / "data" / "seed_decision.json", {"decision": "none"})
    _write_json(root / ".litreview" / "data" / "discovery_campaign.json", {"status": "complete"})
    _write_json(root / ".litreview" / "data" / "landscape.json", {"summary": "Landscape", "streams": []})
    _write_json(root / ".litreview" / "data" / "selected_direction.json", {"title": "Direction"})
    _write_json(root / ".litreview" / "data" / "blueprint.json", {"title": "Blueprint", "sections": []})
    if with_evidence:
        _write_json(root / ".litreview" / "data" / "evidence_map.json", {"summary": "Evidence"})
        _write_json(
            root / ".litreview" / "data" / "fulltext_resolution.json",
            {"coverage_complete": True, "toolkit_oa_full_text_records": 0},
        )
    if with_working_draft:
        _write_json(
            root / ".litreview" / "data" / "working_draft.json",
            {"title": "Draft", "saved_at": "2026-08-29T02:00:00+00:00", "sections": []},
        )


def test_workflow_does_not_force_fulltext_before_first_evidence_map(tmp_path: Path) -> None:
    _workflow_fixture(tmp_path, with_evidence=False, with_working_draft=False)
    result = project_next_step(tmp_path)
    assert result["next_action"] == "construct_evidence_map"
    assert result["skill"] == "litreview-map"


def test_workflow_builds_working_draft_before_handoff(tmp_path: Path) -> None:
    _workflow_fixture(tmp_path, with_evidence=True, with_working_draft=False)
    result = project_next_step(tmp_path)
    assert result["next_action"] == "construct_working_draft"
    assert result["skill"] == "litreview-draft"

    _write_json(
        tmp_path / ".litreview" / "data" / "working_draft.json",
        {"title": "Draft", "saved_at": "2026-08-29T02:00:00+00:00", "sections": []},
    )
    package_step = project_next_step(tmp_path)
    assert package_step["next_action"] == "prepare_researcher_package"
    prepare_researcher_package(tmp_path, export_word=False)
    handoff = project_next_step(tmp_path)
    assert handoff["next_action"] == "researcher_handoff"
    assert any("package prepare" in command for command in handoff["optional_commands"])
