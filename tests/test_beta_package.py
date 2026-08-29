import json
from pathlib import Path

from docx import Document

from litreview_construct.project import init_project
from litreview_construct.researcher_package import canonical_paper_stem, prepare_researcher_package
from litreview_construct.word_export import export_artifact_docx


def _write_jsonl(path: Path, rows: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("".join(json.dumps(row) + "\n" for row in rows), encoding="utf-8")


def _write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value), encoding="utf-8")


def test_init_creates_researcher_facing_library(tmp_path: Path) -> None:
    init_project(tmp_path)
    assert (tmp_path / "papers" / "full_text").is_dir()
    assert (tmp_path / "papers" / "abstract_only").is_dir()
    assert (tmp_path / "papers" / "user_uploads").is_dir()
    assert (tmp_path / "references").is_dir()


def test_doi_filename_is_windows_safe_and_recognizable() -> None:
    row = {"doi": "https://doi.org/10.1016/j.jbankfin.2024.107123", "paper_id": "p1"}
    stem = canonical_paper_stem(row)
    assert stem == "doi_10.1016__j.jbankfin.2024.107123"
    assert "/" not in stem
    assert ":" not in stem


def test_package_exports_working_refs_and_abstract_only_notes(tmp_path: Path) -> None:
    init_project(tmp_path)
    cache = tmp_path / ".litreview" / "cache" / "fulltext"
    cache.mkdir(parents=True, exist_ok=True)
    cached_pdf = cache / "p1.pdf"
    cached_pdf.write_bytes(b"%PDF-1.4\n%%EOF")

    papers = [
        {
            "paper_id": "p1",
            "title": "Bank Efficiency and Ownership",
            "authors": ["Nguyen, An", "Tran, Binh"],
            "year": 2024,
            "journal": "Journal of Banking Research",
            "doi": "10.1016/j.jbankfin.2024.107123",
            "triage_label": "relevant",
            "triage_priority": "core_candidate",
            "file_reference": ".litreview/cache/fulltext/p1.pdf",
            "file_instances": [{"file_reference": ".litreview/cache/fulltext/p1.pdf", "location_type": "managed"}],
            "file_hash": "abc",
            "full_text_provenance": {"access": "open_access", "provider": "openalex"},
        },
        {
            "paper_id": "p2",
            "title": "Governance Heterogeneity in Vietnamese Banks",
            "authors": ["Le, Chi"],
            "year": 2022,
            "journal": "Finance Review",
            "doi": "10.1234/example.2",
            "abstract": "This study reports ownership heterogeneity in the sample.",
            "triage_label": "relevant",
            "triage_priority": "high",
        },
        {
            "paper_id": "p3",
            "title": "Unused Discovery Record",
            "authors": ["Other, Author"],
            "year": 2020,
            "doi": "10.9999/unused",
            "abstract": "Should not be exported because it is not used by the working draft.",
            "triage_label": "relevant",
            "triage_priority": "medium",
        },
    ]
    _write_jsonl(tmp_path / ".litreview" / "data" / "papers.jsonl", papers)
    _write_json(
        tmp_path / ".litreview" / "data" / "working_draft.json",
        {
            "saved_at": "2026-08-29T10:00:00+00:00",
            "sections": [
                {
                    "section_id": "s1",
                    "title": "Ownership and governance",
                    "fragments": [{"paper_ids": ["p1", "p2"]}],
                }
            ],
        },
    )

    result = prepare_researcher_package(tmp_path, export_word=False)

    assert result["working_reference_count"] == 2
    assert result["working_full_text_count"] == 1
    assert result["working_abstract_only_count"] == 1
    assert (tmp_path / "papers" / "full_text" / "doi_10.1016__j.jbankfin.2024.107123.pdf").is_file()
    assert (tmp_path / "papers" / "abstract_only" / "doi_10.1234__example.2.md").is_file()
    assert not (tmp_path / "papers" / "abstract_only" / "doi_10.9999__unused.md").exists()

    enw = (tmp_path / "references" / "references_used.enw").read_text(encoding="utf-8")
    assert "%T Bank Efficiency and Ownership" in enw
    assert "%A Nguyen, An" in enw
    assert "%R 10.1016/j.jbankfin.2024.107123" in enw
    assert "%T Governance Heterogeneity in Vietnamese Banks" in enw
    assert "Unused Discovery Record" not in enw

    csv_text = (tmp_path / "references" / "references_used.csv").read_text(encoding="utf-8")
    assert "Ownership and governance" in csv_text
    assert "pending" in csv_text


def test_word_handoff_is_a_researcher_writing_pack_not_technical_dump(tmp_path: Path) -> None:
    init_project(tmp_path, name="Vietnam Bank Efficiency")
    project_file = tmp_path / ".litreview" / "project.yaml"
    project = project_file.read_text(encoding="utf-8")
    project = project.replace(
        "topic: null",
        "topic: Financial liberalization and operational efficiency of Vietnamese commercial banks",
    ).replace(
        "research_question: null",
        "research_question: How does ownership-governance heterogeneity shape the liberalization-efficiency relationship?",
    )
    project_file.write_text(project, encoding="utf-8")

    _write_json(
        tmp_path / ".litreview" / "data" / "selected_direction.json",
        {
            "title": "Ownership-Governance Heterogeneity in Liberalization-Efficiency Transmission",
            "research_idea": "Examine whether ownership and governance alter the relationship between liberalization and bank efficiency.",
            "possible_gap": "Within the reviewed corpus, direct interaction evidence remains limited.",
        },
    )
    _write_json(
        tmp_path / ".litreview" / "data" / "blueprint.json",
        {
            "organizing_logic": "Move from reform context to efficiency, ownership-governance channels, and the core moderation argument.",
            "sections": [
                {
                    "section_id": "s1",
                    "title": "Ownership and governance channels",
                    "purpose": "Explain the main institutional channel.",
                    "key_arguments": ["Compare state, joint-stock, and foreign-bank heterogeneity."],
                    "theoretical_foundations": ["Quiet-life and competition arguments"],
                    "unresolved_questions": ["How should equitization intensity be measured?"],
                    "transition_logic": "Connect ownership heterogeneity to liberalization-efficiency transmission.",
                    "anchor_paper_ids": ["p1"],
                    "supporting_paper_ids": ["p2"],
                    "conflicting_paper_ids": [],
                }
            ],
            "cross_section_synthesis_tasks": ["Reconcile efficiency and profitability findings."],
            "closing_tasks": ["State only corpus-bounded gaps that survive verification."],
        },
    )
    _write_json(
        tmp_path / ".litreview" / "data" / "working_draft.json",
        {
            "saved_at": "2026-08-29T10:00:00+00:00",
            "sections": [
                {
                    "section_id": "s1",
                    "title": "Ownership and governance channels",
                    "framing_note": "Synthesize ownership-governance heterogeneity without treating abstract evidence as confirmed.",
                    "fragments": [
                        {
                            "draft_text": "Available evidence suggests that ownership and governance may shape efficiency heterogeneity across Vietnamese banks.",
                            "paper_ids": ["p1", "p2"],
                            "researcher_tasks": ["Verify the ownership classification against full texts."],
                            "verification_notes": ["One supporting paper remains abstract-only."],
                        }
                    ],
                    "transition_draft": "The next section links this heterogeneity to liberalization shocks.",
                    "researcher_decisions": ["Confirm whether equitization is treated as intensity rather than a dummy."],
                }
            ],
            "cross_section_notes": ["Keep efficiency distinct from profitability."],
            "final_researcher_tasks": ["Rewrite retained fragments in the researcher's scholarly voice."],
            "limitations": ["Some detailed claims still require full-text verification."],
        },
    )
    _write_jsonl(
        tmp_path / ".litreview" / "data" / "papers.jsonl",
        [
            {
                "paper_id": "p1",
                "title": "Bank Efficiency and Ownership in Vietnam",
                "authors": ["Nguyen, An"],
                "year": 2024,
                "journal": "Journal of Banking Research",
                "doi": "10.1234/core.1",
                "file_reference": "papers/full_text/doi_10.1234__core.1.pdf",
            },
            {
                "paper_id": "p2",
                "title": "Governance Heterogeneity in Vietnamese Banks",
                "authors": ["Le, Chi"],
                "year": 2022,
                "journal": "Finance Review",
                "doi": "10.1234/support.2",
            },
        ],
    )
    _write_jsonl(
        tmp_path / ".litreview" / "data" / "evidence.jsonl",
        [
            {"evidence_id": "e1", "paper_id": "p1", "source_basis": "full_text"},
            {"evidence_id": "e2", "paper_id": "p2", "source_basis": "abstract"},
        ],
    )

    result = export_artifact_docx(tmp_path, artifact="handoff")
    path = Path(result["output"])
    assert path.name == "LitReview_Researcher_Writing_Pack.docx"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    combined = text + "\n" + table_text

    for expected in (
        "Research focus",
        "Literature review plan",
        "Working draft material",
        "Source verification checklist",
        "Final writing checklist",
        "Working references and files",
        "Bank Efficiency and Ownership in Vietnam",
        "AI checked against full text; researcher verification pending",
    ):
        assert expected in combined

    for forbidden in (
        "paper_id",
        "evidence_id",
        ".litreview",
        "OpenAlex",
        "provider_failures",
        "AI-use statement",
        "activity.jsonl",
    ):
        assert forbidden not in combined
