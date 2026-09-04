import json
from pathlib import Path

from docx import Document

from litreview_construct.ai_use import generate_ai_use_statement
from litreview_construct.blueprint import accept_blueprint, save_blueprint
from litreview_construct.corpus import rank_corpus, record_decision
from litreview_construct.direction import apply_direction_decision, save_direction_candidates
from litreview_construct.draft_support import save_working_draft
from litreview_construct.evidence import save_evidence_map
from litreview_construct.intent import accept_intent, set_intent
from litreview_construct.landscape import save_landscape
from litreview_construct.project import init_project
from litreview_construct.researcher_package import prepare_researcher_package
from litreview_construct.seed_state import skip_seed_literature
from litreview_construct.workflow import project_next_step


def _write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value), encoding="utf-8")


def _write_jsonl(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("".join(json.dumps(row) + "\n" for row in rows), encoding="utf-8")


def test_researcher_journey_reaches_researcher_writing_pack_not_final_review(tmp_path: Path) -> None:
    # 1. Researcher-owned Intent and seed decision.
    init_project(tmp_path, name="Working Capital Product Journey")
    set_intent(
        tmp_path,
        topic="working capital management and firm performance",
        research_question="How does working capital management affect firm performance, and when might the relationship be nonlinear?",
        publication_from=2010,
        publication_to=2026,
        languages=["en"],
    )
    accept_intent(tmp_path)
    assert project_next_step(tmp_path)["next_action"] == "ask_seed_literature"
    skip_seed_literature(tmp_path)
    assert project_next_step(tmp_path)["next_action"] == "start_discovery"

    state_root = tmp_path / ".litreview"

    # 2. Durable result of a completed multi-source narrative discovery campaign.
    papers = [
        {
            "paper_id": "p1",
            "title": "Working Capital Management and Firm Performance",
            "authors": ["Author A"],
            "year": 2020,
            "journal": "Finance Journal",
            "doi": "10.1000/p1",
            "abstract": "Working capital management is associated with firm performance.",
            "source_origin": "openalex",
            "discovery_sources": ["openalex", "crossref"],
            "triage_label": "relevant",
            "triage_priority": "core_candidate",
            "status": "relevant",
        },
        {
            "paper_id": "p2",
            "title": "Nonlinear Working Capital and Firm Value",
            "authors": ["Author B"],
            "year": 2023,
            "journal": "Corporate Finance Review",
            "doi": "10.1000/p2",
            "abstract": "The relationship between working capital and firm value may be nonlinear.",
            "source_origin": "semantic_scholar",
            "discovery_sources": ["semantic_scholar"],
            "triage_label": "relevant",
            "triage_priority": "core_candidate",
            "status": "relevant",
        },
        {
            "paper_id": "p3",
            "title": "Financing Constraints and Working Capital Policy",
            "authors": ["Author C"],
            "year": 2022,
            "journal": "Emerging Markets Journal",
            "doi": "10.1000/p3",
            "abstract": "Financing constraints shape working capital policy.",
            "source_origin": "crossref",
            "discovery_sources": ["crossref"],
            "triage_label": "background",
            "triage_priority": "high",
            "status": "background",
        },
    ]
    _write_jsonl(state_root / "data" / "papers.jsonl", papers)
    _write_json(
        state_root / "data" / "discovery_campaign.json",
        {
            "campaign_id": "campaign-e2e",
            "status": "complete",
            "iterations": [
                {"phase": "broad", "queries": ["working capital firm performance"], "providers": ["openalex", "crossref", "semantic_scholar"]},
                {"phase": "focused", "queries": ["nonlinear working capital firm value"], "providers": ["openalex", "semantic_scholar"]},
            ],
            "review_checkpoints": [
                {"checkpoint_id": "review-1", "decision": {"action": "focus"}},
                {"checkpoint_id": "review-2", "decision": {"action": "finish"}},
            ],
            "selected_focuses": ["Nonlinear working-capital optimization"],
            "researcher_completion": {"action": "finish"},
        },
    )
    state_file = state_root / "state.json"
    state = json.loads(state_file.read_text(encoding="utf-8"))
    state["stages"]["literature_discovery"]["status"] = "accepted"
    state["current_stage"] = "literature_discovery"
    _write_json(state_file, state)

    # 3. Post-triage corpus refinement: researcher chooses to narrow before any bulk download.
    retained_step = project_next_step(tmp_path)
    assert retained_step["next_action"] == "retained_corpus_checkpoint"
    assert retained_step["human_checkpoint_required"] is True
    record_decision(tmp_path, stage="retained", action="refine")
    assert project_next_step(tmp_path)["next_action"] == "rank_evidence_candidates"
    assert rank_corpus(tmp_path, to_tier="evidence")["selected_records"] == 3

    evidence_step = project_next_step(tmp_path)
    assert evidence_step["next_action"] == "evidence_candidate_checkpoint"
    record_decision(tmp_path, stage="evidence", action="refine")
    assert project_next_step(tmp_path)["next_action"] == "rank_core_papers"
    assert rank_corpus(tmp_path, to_tier="core")["selected_records"] == 3

    core_step = project_next_step(tmp_path)
    assert core_step["next_action"] == "core_paper_checkpoint"
    record_decision(tmp_path, stage="core", action="continue")
    assert project_next_step(tmp_path)["next_action"] == "construct_current_research_landscape"

    # 4. Research Landscape.
    landscape_file = tmp_path / "landscape_submission.json"
    _write_json(
        landscape_file,
        {
            "summary": "The literature contains a direct performance stream, a nonlinear optimization stream, and financing-constraint context.",
            "anchor_paper_ids": ["p1", "p2"],
            "streams": [
                {
                    "name": "Working-capital performance",
                    "description": "Direct working-capital/performance evidence.",
                    "paper_ids": ["p1", "p3"],
                    "anchor_paper_ids": ["p1"],
                    "main_theories": ["liquidity-profitability trade-off"],
                    "main_methods": ["panel regression"],
                    "major_findings": ["Working-capital choices are associated with firm performance."],
                    "contradictions": [],
                    "recent_developments": [],
                    "confidence": "medium",
                },
                {
                    "name": "Nonlinear working-capital optimization",
                    "description": "Evidence that both insufficient and excessive working capital may be costly.",
                    "paper_ids": ["p2"],
                    "anchor_paper_ids": ["p2"],
                    "main_theories": ["optimal investment trade-off"],
                    "main_methods": ["nonlinear panel specifications"],
                    "major_findings": ["The relationship may be nonlinear."],
                    "contradictions": [],
                    "recent_developments": [],
                    "confidence": "medium",
                },
            ],
            "major_debates": ["Whether the performance relationship is linear or nonlinear."],
            "methodological_clusters": ["Panel regressions", "Nonlinear specifications"],
            "recent_developments": ["Nonlinear working-capital optimization"],
            "unresolved_questions": ["How stable are nonlinear benchmarks across firm contexts?"],
            "limitations": ["Core records still require full-text verification."],
        },
    )
    assert save_landscape(tmp_path, landscape_file)["status"] == "ready_for_review"

    # Full-text acquisition is optional at the corpus checkpoints; continuing does not force it.
    assert project_next_step(tmp_path)["next_action"] == "construct_evidence_map"

    # 5. Evidence Map remains abstract-explicit and provisional.
    evidence_file = tmp_path / "evidence_submission.json"
    _write_json(
        evidence_file,
        {
            "summary": "Abstract-level evidence supports the focal relationship and a possible nonlinear pattern while retaining verification needs.",
            "evidence_items": [
                {
                    "paper_id": "p1",
                    "evidence_type": "association",
                    "claim": "Working capital management is associated with firm performance.",
                    "provenance": "source_reported",
                    "source_basis": "abstract",
                    "source_locator": "abstract",
                    "variables": ["working capital management", "firm performance"],
                    "theories": [],
                    "methods": [],
                    "data_context": [],
                    "certainty": "medium",
                },
                {
                    "paper_id": "p2",
                    "evidence_type": "heterogeneous_finding",
                    "claim": "The working-capital/performance relationship may be nonlinear.",
                    "provenance": "source_reported",
                    "source_basis": "abstract",
                    "source_locator": "abstract",
                    "variables": ["working capital", "firm value"],
                    "theories": [],
                    "methods": [],
                    "data_context": [],
                    "certainty": "medium",
                },
                {
                    "paper_id": "p3",
                    "evidence_type": "theory",
                    "claim": "Financing constraints can shape working-capital policy.",
                    "provenance": "source_reported",
                    "source_basis": "abstract",
                    "source_locator": "abstract",
                    "variables": ["financing constraints", "working capital policy"],
                    "theories": [],
                    "methods": [],
                    "data_context": [],
                    "certainty": "medium",
                },
            ],
            "cross_paper_patterns": ["Working-capital decisions are linked to performance and financing context."],
            "contradictions": [],
            "evidence_gaps": ["Threshold stability remains unverified."],
            "papers_requiring_full_text": ["p1", "p2", "p3"],
            "limitations": ["Current substantive evidence is abstract-based."],
        },
    )
    assert save_evidence_map(tmp_path, evidence_file)["requires_full_text"] == 3
    assert project_next_step(tmp_path)["next_action"] == "propose_research_directions"
    evidence_rows = [json.loads(line) for line in (state_root / "data" / "evidence.jsonl").read_text(encoding="utf-8").splitlines() if line.strip()]
    eids = {row["paper_id"]: row["evidence_id"] for row in evidence_rows}

    # 6. AI proposes; researcher selects.
    directions_file = tmp_path / "directions_submission.json"
    _write_json(
        directions_file,
        {
            "summary": "Two defensible directions remain provisional pending fuller source verification.",
            "directions": [
                {
                    "title": "Nonlinear working-capital optimization",
                    "research_idea": "Test whether an intermediate working-capital range is associated with stronger performance.",
                    "rationale": "The current evidence suggests a nonlinear relationship.",
                    "supporting_paper_ids": ["p1", "p2"],
                    "supporting_evidence_ids": [eids["p1"], eids["p2"]],
                    "what_is_known": ["Working-capital choices are associated with firm performance."],
                    "possible_gap": "Within the reviewed evidence, threshold stability remains insufficiently verified.",
                    "novelty": "Test nonlinear benchmark behavior under a clearly defined firm context.",
                    "data_feasibility": "Potentially feasible with firm accounting data.",
                    "methodological_feasibility": "Nonlinear panel specifications are plausible.",
                    "difficulty": "medium",
                    "risks": ["Benchmarks may be sample-specific."],
                    "limitations": ["Current evidence is abstract-heavy."],
                    "verification_needs": ["Verify nonlinear specifications in full text."],
                    "confidence": "medium",
                },
                {
                    "title": "Financing constraints as a boundary condition",
                    "research_idea": "Test whether financing constraints alter working-capital/performance associations.",
                    "rationale": "Financing constraints shape working-capital policy.",
                    "supporting_paper_ids": ["p1", "p3"],
                    "supporting_evidence_ids": [eids["p1"], eids["p3"]],
                    "what_is_known": ["Financing constraints affect working-capital choices."],
                    "possible_gap": "The moderating relationship is not established by the current bounded evidence.",
                    "novelty": "Integrate financing constraints into the focal relationship.",
                    "data_feasibility": "Potentially feasible if constraints can be proxied.",
                    "methodological_feasibility": "Interaction specifications are plausible.",
                    "difficulty": "medium",
                    "risks": ["Proxy choice may drive inference."],
                    "limitations": ["Current evidence does not establish causality."],
                    "verification_needs": ["Verify measurement approaches."],
                    "confidence": "low",
                },
            ],
            "cross_direction_notes": ["Both require source verification."],
            "limitations": ["The evidence base is abstract-heavy."],
        },
    )
    assert save_direction_candidates(tmp_path, directions_file)["human_decision_required"] is True
    direction_rows = [json.loads(line) for line in (state_root / "data" / "directions.jsonl").read_text(encoding="utf-8").splitlines() if line.strip()]
    decision_file = tmp_path / "direction_decision.json"
    _write_json(
        decision_file,
        {"action": "select", "direction_ids": [direction_rows[0]["direction_id"]], "researcher_notes": "Use the nonlinear direction."},
    )
    assert apply_direction_decision(tmp_path, decision_file)["status"] == "accepted"
    assert project_next_step(tmp_path)["next_action"] == "construct_literature_review_blueprint"

    # 7. Evidence-linked Blueprint and explicit researcher acceptance.
    blueprint_file = tmp_path / "blueprint_submission.json"
    _write_json(
        blueprint_file,
        {
            "title": "Working-capital literature review architecture",
            "organizing_logic": "Move from benchmark logic to nonlinear evidence and financing constraints.",
            "opening_tasks": ["Define the focal working-capital construct."],
            "sections": [
                {
                    "title": "Working-capital benchmark and nonlinear optimality",
                    "purpose": "Establish why both insufficient and excessive working capital may be costly.",
                    "key_arguments": ["Develop the benchmark logic conservatively."],
                    "anchor_paper_ids": ["p1", "p2"],
                    "supporting_paper_ids": [],
                    "conflicting_paper_ids": [],
                    "evidence_ids": [eids["p1"], eids["p2"]],
                    "theoretical_foundations": ["Liquidity-profitability trade-off"],
                    "methodological_context": ["Nonlinear panel specifications"],
                    "hypothesis_or_proposition_links": ["Motivates H1."],
                    "unresolved_questions": ["How should the benchmark be estimated?"],
                    "transition_logic": "Move from benchmark logic to financing constraints.",
                },
                {
                    "title": "Financing constraints and underinvestment exposure",
                    "purpose": "Explain why constrained firms may operate below the feasible benchmark.",
                    "key_arguments": ["Use financing constraints as a plausible mechanism, not a proven causal channel."],
                    "anchor_paper_ids": ["p3"],
                    "supporting_paper_ids": ["p1"],
                    "conflicting_paper_ids": [],
                    "evidence_ids": [eids["p3"], eids["p1"]],
                    "theoretical_foundations": ["Financing constraints"],
                    "methodological_context": [],
                    "hypothesis_or_proposition_links": ["Motivates H2."],
                    "unresolved_questions": ["Which constraint proxy is most defensible?"],
                    "transition_logic": None,
                },
            ],
            "cross_section_synthesis_tasks": ["Reconcile nonlinear and financing-constraint evidence."],
            "closing_tasks": ["State only novelty that survives verification."],
            "verification_priorities": ["Verify all core papers in full text."],
            "limitations": ["The accepted architecture currently includes abstract-based evidence."],
        },
    )
    assert save_blueprint(tmp_path, blueprint_file)["status"] == "ready_for_review"
    assert project_next_step(tmp_path)["next_action"] == "researcher_blueprint_review"
    assert accept_blueprint(tmp_path)["status"] == "accepted"
    assert project_next_step(tmp_path)["next_action"] == "construct_working_draft"

    # 8. Bounded Working Draft, not final manuscript prose.
    blueprint = json.loads((state_root / "data" / "blueprint.json").read_text(encoding="utf-8"))
    draft_sections = []
    for section in blueprint["sections"]:
        paper_ids = list(dict.fromkeys([*section.get("anchor_paper_ids", []), *section.get("supporting_paper_ids", [])]))
        draft_sections.append(
            {
                "section_id": section["section_id"],
                "title": section["title"],
                "framing_note": section["purpose"],
                "fragments": [
                    {
                        "purpose": section["purpose"],
                        "draft_text": "Available evidence suggests the accepted argument while the current associational and verification boundaries remain unresolved.",
                        "paper_ids": paper_ids,
                        "evidence_ids": section.get("evidence_ids", []),
                        "researcher_tasks": ["Verify sources and rewrite in the researcher's scholarly voice."],
                        "verification_notes": ["Evidence remains abstract-based in this deterministic fixture."],
                    }
                ],
                "transition_draft": section.get("transition_logic"),
                "researcher_decisions": [],
            }
        )
    working_file = state_root / "packets" / "working_draft_submission.json"
    _write_json(
        working_file,
        {
            "title": "Researcher Working Draft: Working Capital and Firm Performance",
            "opening_note": "Use this as editable working material, not final manuscript prose.",
            "sections": draft_sections,
            "cross_section_notes": ["Preserve the association-versus-causality boundary."],
            "final_researcher_tasks": ["Obtain full text, verify citations, and rewrite all retained prose."],
            "limitations": ["The present draft is based on abstract-level evidence."],
        },
    )
    saved_draft = save_working_draft(tmp_path, working_file)
    assert saved_draft["sections"] == 2
    assert (tmp_path / "outputs" / "06b_literature_review_working_draft.md").exists()

    # 9. Technical package preparation occurs before the final researcher handoff.
    package_step = project_next_step(tmp_path)
    assert package_step["next_action"] == "prepare_researcher_package"
    package = prepare_researcher_package(tmp_path, export_word=True)
    assert package["working_reference_count"] == 3
    assert (tmp_path / "references" / "references_used.enw").exists()
    writing_pack = Path(package["word_handoff"])
    if not writing_pack.is_absolute():
        writing_pack = tmp_path / writing_pack
    assert writing_pack.name == "LitReview_Researcher_Writing_Pack.docx"
    assert writing_pack.exists()

    handoff = project_next_step(tmp_path)
    assert handoff["next_action"] == "researcher_handoff"
    assert handoff["human_checkpoint_required"] is True
    assert handoff["prohibited_next_step"] == "present_unverified_ai_draft_as_submission_ready_final_review"

    # Word handoff is researcher-facing, not a concatenated technical report.
    document = Document(writing_pack)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    combined = text + "\n" + table_text
    assert "Literature Review — Researcher Writing Pack" in combined
    assert "Literature review plan" in combined
    assert "Working draft material" in combined
    assert "Source verification checklist" in combined
    assert "Working references and files" in combined
    assert "paper_id" not in combined
    assert "evidence_id" not in combined
    assert ".litreview" not in combined
    assert "provider_failures" not in combined

    # AI-use disclosure remains separate and activity-grounded.
    disclosure = generate_ai_use_statement(tmp_path, style="standard")
    assert disclosure["statement"]
    assert (tmp_path / "outputs" / "07_ai_use_statement.md").exists()

    # There is deliberately no complete final submission-ready review artifact.
    assert (tmp_path / "outputs" / "03_research_landscape.md").exists()
    assert (tmp_path / "outputs" / "04_evidence_map.md").exists()
    assert (tmp_path / "outputs" / "05_research_direction.md").exists()
    assert (tmp_path / "outputs" / "06_literature_review_blueprint.md").exists()
    assert (tmp_path / "outputs" / "06b_literature_review_working_draft.md").exists()
    assert not (tmp_path / "outputs" / "08_final_literature_review.md").exists()
