from __future__ import annotations

import json
import re
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

from .activity import append_activity
from .bibliography import normalize_doi
from .project import PROJECT_DIR


ARTIFACTS = {
    "intent": ("outputs/01_research_intent.md", "01_research_intent.docx"),
    "seed-inventory": ("outputs/02_seed_inventory.md", "02_seed_inventory.docx"),
    "landscape": ("outputs/03_research_landscape.md", "03_research_landscape.docx"),
    "evidence": ("outputs/04_evidence_map.md", "04_evidence_map.docx"),
    "direction": ("outputs/05_research_direction.md", "05_research_direction.docx"),
    "blueprint": ("outputs/06_literature_review_blueprint.md", "06_literature_review_blueprint.docx"),
    "working-draft": ("outputs/06b_literature_review_working_draft.md", "06b_literature_review_working_draft.docx"),
    "ai-use": ("outputs/07_ai_use_statement.md", "07_ai_use_statement.docx"),
}


def _project_root(root: Path) -> Path:
    root = root.expanduser().resolve()
    if not (root / PROJECT_DIR / "project.yaml").exists():
        raise FileNotFoundError(f"No Lit Review Construct project found at {root}")
    return root


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    for name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def _split_inline(text: str) -> list[tuple[str, str]]:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")
    parts: list[tuple[str, str]] = []
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            parts.append(("text", text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            parts.append(("bold", token[2:-2]))
        elif token.startswith("`"):
            parts.append(("code", token[1:-1]))
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()  # type: ignore[union-attr]
            parts.append(("link", f"{label} ({url})"))
        pos = match.end()
    if pos < len(text):
        parts.append(("text", text[pos:]))
    return parts


def _add_inline(paragraph, text: str) -> None:
    for kind, value in _split_inline(text):
        run = paragraph.add_run(value)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif kind == "link":
            run.underline = True


def _table_cells(line: str) -> list[str]:
    text = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in text.split("|")]


def _is_separator(line: str) -> bool:
    cells = _table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _add_table(document: Document, lines: list[str]) -> None:
    rows = [_table_cells(line) for line in lines if line.strip()]
    if len(rows) >= 2 and _is_separator(lines[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index in range(width):
            text = row[c_index] if c_index < len(row) else ""
            paragraph = table.cell(r_index, c_index).paragraphs[0]
            _add_inline(paragraph, text)
            if r_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def _append_markdown(document: Document, markdown: str, *, suppress_first_title: bool = False) -> None:
    lines = markdown.splitlines()
    i = 0
    first_heading_seen = False
    while i < len(lines):
        raw = lines[i].rstrip()
        if raw.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|") and _is_separator(lines[i + 1]):
            table_lines = [raw, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            _add_table(document, table_lines)
            continue
        if not raw.strip():
            i += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", raw)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if suppress_first_title and not first_heading_seen and level == 1:
                first_heading_seen = True
                i += 1
                continue
            first_heading_seen = True
            document.add_heading(text, level=min(level, 3))
            i += 1
            continue
        if raw.startswith(">"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(raw.lstrip("> "))
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", raw)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline(paragraph, numbered.group(2))
            i += 1
            continue
        if raw.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, raw[2:])
            i += 1
            continue
        paragraph = document.add_paragraph()
        _add_inline(paragraph, raw)
        i += 1


def _project_name(root: Path) -> str:
    payload = yaml.safe_load((root / PROJECT_DIR / "project.yaml").read_text(encoding="utf-8")) or {}
    return str(payload.get("name") or payload.get("project_name") or root.name)


def _load_json(path: Path) -> dict[str, object]:
    if not path.exists():
        return {}
    value = json.loads(path.read_text(encoding="utf-8"))
    return value if isinstance(value, dict) else {}


def _load_jsonl(path: Path) -> list[dict[str, object]]:
    if not path.exists():
        return []
    rows: list[dict[str, object]] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        if not raw.strip():
            continue
        try:
            value = json.loads(raw)
        except json.JSONDecodeError:
            continue
        if isinstance(value, dict):
            rows.append(value)
    return rows


def _add_bullets(document: Document, values: list[object]) -> None:
    for value in values:
        text = str(value).strip()
        if text:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, text)


def _human_evidence_status(
    paper_id: str,
    paper: dict[str, object],
    evidence_by_paper: dict[str, list[dict[str, object]]],
) -> tuple[str, str]:
    local = bool(paper.get("file_reference") or paper.get("file_hash"))
    evidence = evidence_by_paper.get(paper_id) or []
    ai_full_text = any(str(row.get("source_basis") or "") == "full_text" for row in evidence)
    if ai_full_text:
        return "Full text available", "AI checked against full text; researcher verification pending"
    if local:
        return "Full text available", "AI has not yet checked the relevant claim against full text; researcher verification pending"
    return "Abstract/metadata only", "Full-text verification still needed before detailed claims are treated as established"


def _used_paper_ids(working: dict[str, object], blueprint: dict[str, object]) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()

    def add(values: object) -> None:
        if not isinstance(values, list):
            return
        for value in values:
            paper_id = str(value or "")
            if paper_id and paper_id not in seen:
                seen.add(paper_id)
                ordered.append(paper_id)

    for section in working.get("sections") or []:
        if not isinstance(section, dict):
            continue
        for fragment in section.get("fragments") or []:
            if isinstance(fragment, dict):
                add(fragment.get("paper_ids"))
    if ordered:
        return ordered
    for section in blueprint.get("sections") or []:
        if not isinstance(section, dict):
            continue
        add(section.get("anchor_paper_ids"))
        add(section.get("supporting_paper_ids"))
        add(section.get("conflicting_paper_ids"))
    return ordered


def _add_research_focus(document: Document, root: Path, direction: dict[str, object]) -> None:
    project = yaml.safe_load((root / PROJECT_DIR / "project.yaml").read_text(encoding="utf-8")) or {}
    research = project.get("research") or {}
    document.add_heading("1. Research focus", level=1)
    topic = research.get("topic") if isinstance(research, dict) else None
    question = research.get("research_question") if isinstance(research, dict) else None
    if topic:
        p = document.add_paragraph()
        _add_inline(p, f"**Topic:** {topic}")
    if question:
        p = document.add_paragraph()
        _add_inline(p, f"**Research question:** {question}")
    if direction.get("title"):
        p = document.add_paragraph()
        _add_inline(p, f"**Selected research direction:** {direction['title']}")
    if direction.get("research_idea"):
        document.add_paragraph(str(direction["research_idea"]))
    if direction.get("possible_gap"):
        p = document.add_paragraph()
        _add_inline(p, "**Working gap to verify:** " + str(direction["possible_gap"]))


def _add_blueprint_plan(document: Document, blueprint: dict[str, object]) -> None:
    document.add_heading("2. Literature review plan", level=1)
    document.add_paragraph(
        "Use this structure as the argument path for the literature review. The headings and transitions are editable; the purpose is to keep the review coherent while you verify and rewrite the evidence."
    )
    if blueprint.get("organizing_logic"):
        p = document.add_paragraph()
        _add_inline(p, "**Organizing logic:** " + str(blueprint["organizing_logic"]))
    for order, section in enumerate(blueprint.get("sections") or [], start=1):
        if not isinstance(section, dict):
            continue
        document.add_heading(f"{order}. {section.get('title') or 'Untitled section'}", level=2)
        if section.get("purpose"):
            p = document.add_paragraph()
            _add_inline(p, "**Purpose:** " + str(section["purpose"]))
        arguments = list(section.get("key_arguments") or [])
        if arguments:
            p = document.add_paragraph()
            _add_inline(p, "**Points this section should establish:**")
            _add_bullets(document, arguments)
        foundations = list(section.get("theoretical_foundations") or [])
        if foundations:
            p = document.add_paragraph()
            _add_inline(p, "**Useful theoretical foundations:**")
            _add_bullets(document, foundations)
        unresolved = list(section.get("unresolved_questions") or [])
        if unresolved:
            p = document.add_paragraph()
            _add_inline(p, "**Questions you still need to resolve:**")
            _add_bullets(document, unresolved)
        if section.get("transition_logic"):
            p = document.add_paragraph()
            _add_inline(p, "**Transition to the next section:** " + str(section["transition_logic"]))


def _add_working_draft(document: Document, working: dict[str, object]) -> None:
    document.add_heading("3. Working draft material", level=1)
    document.add_paragraph(
        "The text below is working material to revise, verify, shorten, expand, or reject. It is intentionally not a seamless final literature review."
    )
    for order, section in enumerate(working.get("sections") or [], start=1):
        if not isinstance(section, dict):
            continue
        document.add_heading(f"{order}. {section.get('title') or 'Untitled section'}", level=2)
        if section.get("framing_note"):
            p = document.add_paragraph()
            _add_inline(p, "**What this section is trying to do:** " + str(section["framing_note"]))
        for fragment in section.get("fragments") or []:
            if not isinstance(fragment, dict):
                continue
            draft_text = str(fragment.get("draft_text") or "").strip()
            if draft_text:
                document.add_paragraph(draft_text)
            tasks = list(fragment.get("researcher_tasks") or [])
            notes = list(fragment.get("verification_notes") or [])
            decisions = list(section.get("researcher_decisions") or [])
            if tasks or notes or decisions:
                p = document.add_paragraph()
                _add_inline(p, "**Before treating this section as finished:**")
                _add_bullets(document, [*tasks, *notes, *decisions])
        if section.get("transition_draft"):
            p = document.add_paragraph()
            _add_inline(p, "**Possible transition:** " + str(section["transition_draft"]))


def _add_verification_checklist(
    document: Document,
    used_ids: list[str],
    papers_by_id: dict[str, dict[str, object]],
    evidence_by_paper: dict[str, list[dict[str, object]]],
) -> None:
    document.add_heading("4. Source verification checklist", level=1)
    document.add_paragraph(
        "This is the working source set behind the current draft. Verify the most important papers before converting provisional statements into firm claims."
    )
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Paper", "Year", "DOI", "Access", "What remains to verify"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        for run in table.cell(0, index).paragraphs[0].runs:
            run.bold = True
    for paper_id in used_ids:
        paper = papers_by_id.get(paper_id)
        if not paper:
            continue
        access, verification = _human_evidence_status(paper_id, paper, evidence_by_paper)
        row = table.add_row().cells
        row[0].text = str(paper.get("title") or "Untitled")
        row[1].text = str(paper.get("year") or "")
        doi = normalize_doi(str(paper.get("doi"))) if paper.get("doi") else None
        row[2].text = doi or ""
        row[3].text = access
        row[4].text = verification


def _add_final_writing_checklist(document: Document, blueprint: dict[str, object], working: dict[str, object]) -> None:
    document.add_heading("5. Final writing checklist", level=1)
    tasks: list[object] = []
    tasks.extend(blueprint.get("cross_section_synthesis_tasks") or [])
    tasks.extend(blueprint.get("closing_tasks") or [])
    tasks.extend(working.get("cross_section_notes") or [])
    tasks.extend(working.get("final_researcher_tasks") or [])
    seen: set[str] = set()
    deduped: list[object] = []
    for task in tasks:
        text = str(task).strip()
        if text and text not in seen:
            seen.add(text)
            deduped.append(text)
    if deduped:
        _add_bullets(document, deduped)
    else:
        _add_bullets(
            document,
            [
                "Verify core claims against the strongest available full texts.",
                "Resolve any remaining interpretation or construct-definition decisions.",
                "Rewrite the working fragments into your own scholarly voice and connect sections smoothly.",
                "Check every in-text citation against the canonical reference record before submission.",
            ],
        )
    limitations = list(working.get("limitations") or [])
    if limitations:
        p = document.add_paragraph()
        _add_inline(p, "**Current limitations to keep in mind while writing:**")
        _add_bullets(document, limitations)


def _add_reference_help(
    document: Document,
    used_ids: list[str],
    papers_by_id: dict[str, dict[str, object]],
) -> None:
    document.add_heading("6. Working references and files", level=1)
    document.add_paragraph(
        "The EndNote import file for the references currently used by this Blueprint/Working Draft is saved at `references/references_used.enw`. A CSV audit copy is saved beside it."
    )
    document.add_paragraph(
        "Downloaded lawful full texts are in `papers/full_text/`; references still available only at abstract/metadata level are listed in `papers/abstract_only/`; your own PDFs can be placed in `papers/user_uploads/`."
    )
    document.add_heading("Working reference list", level=2)
    for paper_id in used_ids:
        paper = papers_by_id.get(paper_id)
        if not paper:
            continue
        authors = paper.get("authors") if isinstance(paper.get("authors"), list) else []
        author_text = ", ".join(str(value) for value in authors if str(value).strip()) or "Unknown author"
        year = str(paper.get("year") or "n.d.")
        title = str(paper.get("title") or "Untitled")
        journal = str(paper.get("journal") or paper.get("venue") or "").strip()
        doi = normalize_doi(str(paper.get("doi"))) if paper.get("doi") else None
        parts = [f"{author_text} ({year}). {title}."]
        if journal:
            parts.append(journal + ".")
        if doi:
            parts.append(f"https://doi.org/{doi}")
        document.add_paragraph(" ".join(parts), style="List Bullet")


def export_artifact_docx(root: Path, *, artifact: str, output: Path | None = None) -> dict[str, object]:
    root = _project_root(root)
    if artifact == "handoff":
        return export_handoff_docx(root, output=output)
    if artifact not in ARTIFACTS:
        raise ValueError("Unknown Word artifact. Choose: " + ", ".join([*ARTIFACTS, "handoff"]))
    source_rel, default_name = ARTIFACTS[artifact]
    source = root / source_rel
    if not source.exists():
        raise ValueError(f"Artifact is not available yet: {source_rel}")
    target = output.expanduser().resolve() if output else root / "outputs" / default_name
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure(document)
    _append_markdown(document, source.read_text(encoding="utf-8"))
    document.core_properties.title = f"{_project_name(root)} — {artifact}"
    document.core_properties.subject = "Lit Review Construct researcher artifact"
    document.save(target)
    append_activity(
        root,
        category="document_export",
        actor="toolkit",
        inputs={"format": "docx", "artifact": artifact},
        outputs=[str(target.relative_to(root)) if target.is_relative_to(root) else str(target)],
        notes="Exported a project artifact to Word. Structured project state remains local.",
    )
    return {"artifact": artifact, "source": str(source), "output": str(target)}


def export_handoff_docx(root: Path, *, output: Path | None = None) -> dict[str, object]:
    """Export a researcher writing pack, not a concatenated technical workflow report."""
    root = _project_root(root)
    data_root = root / PROJECT_DIR / "data"
    direction = _load_json(data_root / "selected_direction.json")
    blueprint = _load_json(data_root / "blueprint.json")
    working = _load_json(data_root / "working_draft.json")
    if not blueprint or not working:
        raise ValueError("The accepted Blueprint and Working Draft are required before researcher handoff export.")

    papers = _load_jsonl(data_root / "papers.jsonl")
    evidence = _load_jsonl(data_root / "evidence.jsonl")
    papers_by_id = {str(row.get("paper_id")): row for row in papers if row.get("paper_id")}
    evidence_by_paper: dict[str, list[dict[str, object]]] = {}
    for row in evidence:
        paper_id = str(row.get("paper_id") or "")
        if paper_id:
            evidence_by_paper.setdefault(paper_id, []).append(row)
    used_ids = _used_paper_ids(working, blueprint)

    target = output.expanduser().resolve() if output else root / "outputs" / "LitReview_Researcher_Writing_Pack.docx"
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure(document)
    document.add_heading("Literature Review — Researcher Writing Pack", 0)
    p = document.add_paragraph()
    _add_inline(p, _project_name(root))
    document.add_paragraph(
        "This document is designed to help you finish the literature review. It brings together the accepted research direction, the review structure, researcher-editable draft material, source-verification needs, and the working references. It deliberately omits runtime logs, internal IDs, provider diagnostics, and other implementation detail."
    )
    p = document.add_paragraph()
    run = p.add_run(
        "Use this as a writing and verification pack, not as submission-ready prose. You remain responsible for source verification, citation selection, interpretation, final wording, and authorship."
    )
    run.italic = True

    _add_research_focus(document, root, direction)
    _add_blueprint_plan(document, blueprint)
    _add_working_draft(document, working)
    _add_verification_checklist(document, used_ids, papers_by_id, evidence_by_paper)
    _add_final_writing_checklist(document, blueprint, working)
    _add_reference_help(document, used_ids, papers_by_id)

    document.core_properties.title = f"{_project_name(root)} — Researcher Writing Pack"
    document.core_properties.subject = "Researcher-facing literature review writing and verification pack"
    document.save(target)
    append_activity(
        root,
        category="document_export",
        actor="toolkit",
        inputs={
            "format": "docx",
            "artifact": "handoff",
            "presentation": "researcher_writing_pack",
            "working_references": len(used_ids),
        },
        outputs=[str(target.relative_to(root)) if target.is_relative_to(root) else str(target)],
        source_ids=used_ids,
        notes="Exported a researcher-facing writing pack. Technical workflow logs and internal identifiers are intentionally excluded from the Word handoff.",
    )
    return {
        "artifact": "handoff",
        "presentation": "researcher_writing_pack",
        "working_references": len(used_ids),
        "output": str(target),
    }
